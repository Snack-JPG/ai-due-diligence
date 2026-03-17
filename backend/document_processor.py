from __future__ import annotations

import mimetypes
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
from chromadb.config import Settings as ChromaSettings
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from openpyxl import load_workbook
from pypdf import PdfReader

from backend.config import settings
from backend.database import db
from backend.schemas import ProgressEvent
from datetime import datetime, timezone


@dataclass
class ProcessedDocument:
    document_id: str
    filename: str
    metadata: dict[str, Any]
    chunks_indexed: int


def now() -> datetime:
    return datetime.now(timezone.utc)


def _build_vectorstore(analysis_id: str) -> Chroma:
    embeddings = OpenAIEmbeddings(model=settings.openai_embedding_model, api_key=settings.openai_api_key)
    return Chroma(
        collection_name=f"analysis_{analysis_id}",
        embedding_function=embeddings,
        persist_directory=str(settings.chroma_persist_dir),
        client_settings=ChromaSettings(anonymized_telemetry=False),
    )


def infer_content_type(filename: str) -> str:
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or "application/octet-stream"


def sanitize_filename(name: str) -> str:
    safe = "".join(ch for ch in name if ch.isalnum() or ch in {".", "_", "-"}).strip("._")
    return safe or "document"


def extract_pdf(path: Path) -> tuple[list[Document], int]:
    reader = PdfReader(str(path))
    docs: list[Document] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(page_content=text, metadata={"locator": f"Page {page_number}", "page_number": page_number}))
    return docs, len(reader.pages)


def extract_docx(path: Path) -> tuple[list[Document], int | None]:
    doc = DocxDocument(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    text = "\n\n".join(paragraphs)
    return [Document(page_content=text, metadata={"locator": "Paragraphs"})], None


def extract_xlsx(path: Path) -> tuple[list[Document], int | None]:
    workbook = load_workbook(filename=str(path), data_only=True)
    docs: list[Document] = []
    for sheet in workbook.worksheets:
        rows = []
        for idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if not any(cell is not None and str(cell).strip() for cell in row):
                continue
            serialized = " | ".join("" if cell is None else str(cell) for cell in row)
            rows.append(f"Row {idx}: {serialized}")
        if rows:
            docs.append(Document(page_content="\n".join(rows), metadata={"locator": f"Sheet {sheet.title}", "sheet": sheet.title}))
    return docs, len(workbook.sheetnames)


def extract_csv(path: Path) -> tuple[list[Document], int | None]:
    df = pd.read_csv(path)
    lines = [", ".join(map(str, df.columns.tolist()))]
    for index, row in df.iterrows():
        lines.append(f"Row {index + 1}: " + " | ".join(map(str, row.fillna("").tolist())))
    return [Document(page_content="\n".join(lines), metadata={"locator": "CSV rows"})], len(df.index)


def extract_txt(path: Path) -> tuple[list[Document], int | None]:
    return [Document(page_content=path.read_text(encoding="utf-8", errors="ignore"), metadata={"locator": "Text"})], None


def load_documents(path: Path) -> tuple[list[Document], int | None]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".xlsx":
        return extract_xlsx(path)
    if suffix == ".csv":
        return extract_csv(path)
    if suffix == ".txt":
        return extract_txt(path)
    raise ValueError(f"Unsupported file type: {suffix}")


class DocumentProcessor:
    def __init__(self) -> None:
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )

    def process_analysis(self, analysis_id: str) -> list[ProcessedDocument]:
        vectorstore = _build_vectorstore(analysis_id)
        processed: list[ProcessedDocument] = []
        for doc_row in db.get_documents(analysis_id):
            document_id = doc_row["id"]
            path = Path(doc_row["storage_path"])
            db.add_event(
                ProgressEvent(
                    type="document.processing.started",
                    analysis_id=analysis_id,
                    timestamp=now(),
                    payload={"document_id": document_id, "filename": doc_row["filename"]},
                )
            )
            try:
                raw_docs, page_count = load_documents(path)
                enriched_docs = []
                for item in raw_docs:
                    metadata = {
                        **item.metadata,
                        "analysis_id": analysis_id,
                        "document_id": document_id,
                        "filename": doc_row["filename"],
                        "document_type": path.suffix.lower().lstrip("."),
                    }
                    enriched_docs.append(Document(page_content=item.page_content, metadata=metadata))
                chunks = self.splitter.split_documents(enriched_docs)
                for index, chunk in enumerate(chunks):
                    chunk.metadata["chunk_index"] = index
                if chunks:
                    vectorstore.add_documents(chunks)
                db.update_document_parse_status(document_id, "completed", page_count)
                db.add_event(
                    ProgressEvent(
                        type="document.processing.completed",
                        analysis_id=analysis_id,
                        timestamp=now(),
                        payload={
                            "document_id": document_id,
                            "filename": doc_row["filename"],
                            "page_count": page_count,
                            "chunks_indexed": len(chunks),
                        },
                    )
                )
                processed.append(
                    ProcessedDocument(
                        document_id=document_id,
                        filename=doc_row["filename"],
                        metadata={"page_count": page_count},
                        chunks_indexed=len(chunks),
                    )
                )
            except Exception as exc:
                db.update_document_parse_status(document_id, "failed")
                db.add_event(
                    ProgressEvent(
                        type="document.processing.completed",
                        analysis_id=analysis_id,
                        timestamp=now(),
                        payload={
                            "document_id": document_id,
                            "filename": doc_row["filename"],
                            "parse_status": "failed",
                            "error": str(exc),
                        },
                    )
                )
        if not processed:
            raise ValueError("No uploaded documents could be parsed successfully")
        return processed

    def as_retriever(self, analysis_id: str):
        return _build_vectorstore(analysis_id).as_retriever(search_kwargs={"k": settings.top_k_retrieval})


document_processor = DocumentProcessor()
